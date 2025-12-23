import os
import pandas as pd
from datetime import datetime
from sqlalchemy import create_engine
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT

# ✅ Connexion PostgreSQL via Render
DATABASE_URL = os.getenv("DATABASE_URL")  # Render définit cette variable automatiquement
engine = create_engine(DATABASE_URL)

def appliquer_style_entete(cell):
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')
    tcPr.append(shd)

# ✅ Traduction des mois français vers anglais
def traduire_mois_en_anglais(mois_francais):
    mois_map = {
        "Janvier": "January", "Février": "February", "Mars": "March", "Avril": "April",
        "Mai": "May", "Juin": "June", "Juillet": "July", "Août": "August",
        "Septembre": "September", "Octobre": "October", "Novembre": "November", "Décembre": "December"
    }
    parts = mois_francais.strip().split()
    if len(parts) != 2:
        raise ValueError(f"Format invalide pour le mois sélectionné : {mois_francais}")
    nom_mois, annee = parts
    mois_anglais = mois_map.get(nom_mois)
    if not mois_anglais:
        raise ValueError(f"Mois non reconnu : {nom_mois}")
    return f"{mois_anglais} {annee}"

def format_montant(valeur):
    return f"{round(valeur):,}".replace(",", " ")
def generer_memo_mensuel(mois_selectionne, afficher=False):
    mois_en_anglais = traduire_mois_en_anglais(mois_selectionne)
    mois_datetime = datetime.strptime(mois_en_anglais, "%B %Y")
    mois_debut = mois_datetime.replace(day=1)
    mois_fin = mois_datetime.replace(day=28) + pd.DateOffset(days=4)
    mois_fin = mois_fin - pd.DateOffset(days=mois_fin.day)

    # 📥 Requêtes sur PostgreSQL
    df_liv = pd.read_sql("""
        SELECT id AS livraison_id, date, transporteur_id, depot AS site_id
        FROM livraison
        WHERE date BETWEEN %(deb)s AND %(fin)s
    """, engine, params={"deb": str(mois_debut.date()), "fin": str(mois_fin.date())})

    if df_liv.empty:
        return (pd.DataFrame(), pd.DataFrame()) if afficher else f"Aucune donnée disponible pour {mois_selectionne}"

    df_sites = pd.read_sql("SELECT id AS site_id, nom_site, numero_compte FROM sites", engine)
    df_comp = pd.read_sql("""
        SELECT livraison_id, produit_id AS produit, volume_manquant, commentaire
        FROM compartiments
        WHERE commentaire = 'Remboursable'
    """, engine)
    df_prix = pd.read_sql("SELECT produit_id, prix, date_debut FROM prix_vente", engine)
    df_prix["date_debut"] = pd.to_datetime(df_prix["date_debut"])
    df_trans = pd.read_sql("SELECT id AS transporteur_id, nom FROM transporteurs", engine)

    # 🧠 Fusion des données
    df = pd.merge(df_comp, df_liv, on="livraison_id")
    df = pd.merge(df, df_sites, on="site_id", how="left")
    df["date"] = pd.to_datetime(df["date"])
    df = df[df["date"].between(mois_debut, mois_fin)]
    if df.empty:
        return (pd.DataFrame(), pd.DataFrame()) if afficher else f"Aucune livraison remboursable pour {mois_selectionne}"

    def get_prix(produit_id, date_livraison):
        prix_liste = df_prix[(df_prix["produit_id"] == produit_id) & (df_prix["date_debut"] <= date_livraison)]
        if prix_liste.empty:
            return None
        return prix_liste.sort_values("date_debut", ascending=False).iloc[0]["prix"]

    df["montant"] = df.apply(
        lambda row: float(row["volume_manquant"]) * float(get_prix(row["produit"], row["date"]) or 0),
        axis=1
    )

    df_transpo = df.groupby("transporteur_id")["montant"].sum().reset_index()
    df_transpo = pd.merge(df_transpo, df_trans, on="transporteur_id")
    df_transpo["montant"] = df_transpo["montant"].apply(format_montant)

    df_sites = df.groupby(["nom_site", "numero_compte"])["montant"].sum().reset_index()
    df_sites["montant"] = df_sites["montant"].apply(format_montant)

    # 📄 Création du document Word
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    date_du_jour = datetime.today().strftime("%d-%m-%Y")
    doc.add_paragraph(f"Abidjan, le {date_du_jour}")
    doc.add_paragraph("De : Eustache YAVO")
    doc.add_paragraph("A : Stéphane KONAN")
    doc.add_paragraph("CC : Fatim KANATE, Brigitte ADON")
    doc.add_paragraph(f"Objet : refacturation des manquants sur livraison {mois_selectionne}")

    # … (tables transporteurs et sites comme dans ta version précédente)

    nom_fichier = f"manquants hors freinte RETAIL-B2B {mois_selectionne}.docx"
    doc.save(nom_fichier)

    return (df_transpo, df_sites) if afficher else nom_fichier